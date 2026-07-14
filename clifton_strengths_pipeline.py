"""
CliftonStrengths 34 -> Spreadsheet Pipeline
2b Limitless Internship | AI & Automation

Extracts each person's rank-ordered 34 CliftonStrengths themes and compiles
everyone into a single spreadsheet, replacing manual retyping of 34 data
points per person.

Accepts THREE input formats, since it's unconfirmed which one the team
actually has for each person — drop any mix of these into one folder:

  1. PDF   — the Gallup emailed report (e.g. "Kucukerdogan-Ipek-ALL_34.pdf")
  2. XLSX/CSV — a spreadsheet export or manually-built tracker. Handles:
       - 2b Limitless's actual format: one column per theme name ("Strategic",
         "Competition", ... "Includer"), one row per person, cell = that
         person's rank number (1-34) for that theme — this is what the team
         manually types in from the PDF today.
       - "long" format: one row per theme, with Name/Theme/Rank-ish columns
       - "wide" format: one row per person, 34 columns labelled "Rank 1".."Rank 34",
         each holding a theme name
       - a bare single column of 34 theme names (one person, name = filename)
  3. TXT   — pasted text copied straight from the Gallup portal page.
       Put one person's pasted results per .txt file. Works whether the
       paste kept the "1. Strategic" numbering or is just 34 theme names
       on separate lines in rank order.

USAGE
-----
    python3 clifton_strengths_pipeline.py --input "path/to/reports_folder" --output "CliftonStrengths_Data.xlsx"
    python3 clifton_strengths_pipeline.py --input "path/to/single_report.pdf" --output "CliftonStrengths_Data.xlsx"

Re-running on the same output file merges in new/updated people (matched by
name) and leaves everyone else untouched, so you can keep dropping files in
the folder over time and re-run.

OUTPUT — TWO SEPARATE FILES
----------------------------
1. The data workbook (--output), laid out exactly like the team's own tracker:
     - "CliftonStrengths Data"     — domain banner row, then 34 theme-name columns,
                                      one row per person, cell = that person's RANK
                                      NUMBER (1-34) for that theme
     - "Domain Summary (Top 10)"   — top-10 theme counts per domain, leading domain
2. The chart workbook (--charts-output, defaults to "<output>_Wheels.xlsx"):
     - "Strengths Wheels"          — a domain-quadrant radar chart per person, 2 per
                                      row, each with a name banner. A manual page
                                      break is inserted after every row of charts so
                                      one never gets sliced across a printed page.
   Skip this file entirely with --no-charts for faster runs on large batches.

REQUIREMENTS
------------
pip install pdfplumber openpyxl pandas matplotlib --break-system-packages
"""

import argparse
import re
import sys
import tempfile
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import pdfplumber
from openpyxl import Workbook, load_workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.pagebreak import Break
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Emu

# ---------------------------------------------------------------------------
# Fixed Gallup CliftonStrengths 34 -> Domain mapping (does not vary by person)
# ---------------------------------------------------------------------------
DOMAIN_MAP = {
    "Achiever": "Executing", "Arranger": "Executing", "Belief": "Executing",
    "Consistency": "Executing", "Deliberative": "Executing", "Discipline": "Executing",
    "Focus": "Executing", "Responsibility": "Executing", "Restorative": "Executing",
    "Activator": "Influencing", "Command": "Influencing", "Communication": "Influencing",
    "Competition": "Influencing", "Maximiser": "Influencing", "Maximizer": "Influencing",
    "Self-Assurance": "Influencing", "Significance": "Influencing", "Woo": "Influencing",
    "Adaptability": "Relationship Building", "Connectedness": "Relationship Building",
    "Developer": "Relationship Building", "Empathy": "Relationship Building",
    "Harmony": "Relationship Building", "Includer": "Relationship Building",
    "Individualisation": "Relationship Building", "Individualization": "Relationship Building",
    "Positivity": "Relationship Building", "Relator": "Relationship Building",
    "Analytical": "Strategic Thinking", "Context": "Strategic Thinking",
    "Futuristic": "Strategic Thinking", "Ideation": "Strategic Thinking",
    "Input": "Strategic Thinking", "Intellection": "Strategic Thinking",
    "Learner": "Strategic Thinking", "Strategic": "Strategic Thinking",
}
DOMAIN_ORDER = ["Executing", "Influencing", "Relationship Building", "Strategic Thinking"]

# Canonical column order used for BOTH the spreadsheet and the wheel chart: themes
# grouped by domain (in DOMAIN_ORDER), alphabetical within each domain — this is the
# exact column order the team's own tracker uses.
DOMAIN_THEMES_ORDERED = {d: [] for d in DOMAIN_ORDER}
for _theme, _domain in DOMAIN_MAP.items():
    if _theme in ("Maximizer", "Individualization"):  # skip US-spelling duplicates
        continue
    DOMAIN_THEMES_ORDERED[_domain].append(_theme)
for _d in DOMAIN_THEMES_ORDERED:
    DOMAIN_THEMES_ORDERED[_d].sort()

ORDERED_THEMES = [theme for d in DOMAIN_ORDER for theme in DOMAIN_THEMES_ORDERED[d]]

# Sort longest-first so "Self-Assurance" matches before a stray "Self" would.
_THEME_ALTERNATION = "|".join(sorted(DOMAIN_MAP.keys(), key=len, reverse=True))
_RANK_PATTERN = re.compile(rf'\b(\d{{1,2}})\.\s*({_THEME_ALTERNATION})\b')
_BARE_THEME_PATTERN = re.compile(rf'^({_THEME_ALTERNATION})$', re.IGNORECASE)


def _normalise_theme(raw: str) -> str | None:
    """Match a loose string (extra spaces, punctuation, wrong case) to a canonical theme name."""
    cleaned = re.sub(r'[^A-Za-z\- ]', '', str(raw)).strip()
    for theme in DOMAIN_MAP:
        if cleaned.lower() == theme.lower():
            return theme
    return None


# ---------------------------------------------------------------------------
# Format 1: PDF (Gallup emailed report)
# ---------------------------------------------------------------------------
def extract_from_pdf(pdf_path: Path) -> dict:
    with pdfplumber.open(pdf_path) as pdf:
        page1_text = pdf.pages[0].extract_text() or ""

    header_line = page1_text.splitlines()[0]
    if "|" in header_line:
        name, date = [part.strip() for part in header_line.split("|", 1)]
    else:
        name, date = header_line.strip(), ""

    ranks = {}
    for num, theme in _RANK_PATTERN.findall(page1_text):
        n = int(num)
        if 1 <= n <= 34 and n not in ranks:
            ranks[n] = theme.strip()

    _require_complete(ranks, pdf_path.name)
    return {"name": name, "date": date, "ranks": ranks, "source_file": pdf_path.name}


# ---------------------------------------------------------------------------
# Format 2: pasted text from the Gallup portal (.txt, one person per file)
# ---------------------------------------------------------------------------
def extract_from_text(txt_path: Path) -> dict:
    text = txt_path.read_text(encoding="utf-8", errors="ignore")
    lines = [l.strip() for l in text.splitlines() if l.strip()]

    name = txt_path.stem.replace("_", " ").replace("-", " ").strip()
    date = ""
    # If the first line looks like "NAME | DATE" or "NAME - DATE", use it and drop it.
    if lines and ("|" in lines[0] or re.search(r'\d{1,2}[-/]\d{1,2}[-/]\d{2,4}', lines[0])):
        header = lines[0]
        if "|" in header:
            parts = [p.strip() for p in header.split("|", 1)]
            name, date = parts[0], parts[1] if len(parts) > 1 else ""
        lines = lines[1:]

    numbered_text = "\n".join(lines)
    ranks = {}
    for num, theme in _RANK_PATTERN.findall(numbered_text):
        n = int(num)
        if 1 <= n <= 34 and n not in ranks:
            ranks[n] = theme.strip()

    # Fallback: no "N. Theme" numbering present — assume each line is one theme,
    # already in rank order (this is what a plain copy-paste of a portal list looks like).
    if len(ranks) < 34:
        ranks = {}
        rank_n = 1
        for line in lines:
            theme = _normalise_theme(line)
            if theme is None:
                # allow trailing junk on the line, e.g. "1 Strategic" already handled above;
                # try stripping a leading rank number/bullet before giving up
                stripped = re.sub(r'^[\d]{1,2}[\.\)]?\s*', '', line)
                theme = _normalise_theme(stripped)
            if theme:
                ranks[rank_n] = theme
                rank_n += 1
            if rank_n > 34:
                break

    _require_complete(ranks, txt_path.name)
    return {"name": name, "date": date, "ranks": ranks, "source_file": txt_path.name}


# ---------------------------------------------------------------------------
# Format 3: Excel / CSV export from the Gallup portal
# ---------------------------------------------------------------------------
def _load_raw_grid(path: Path) -> list:
    """Every cell value, row by row, with NO assumption about which row is the header."""
    if path.suffix.lower() == ".csv":
        import csv
        with open(path, newline="", encoding="utf-8-sig") as f:
            return [row for row in csv.reader(f)]
    wb = load_workbook(path, data_only=True)
    ws = wb.active
    return [[c for c in row] for row in ws.iter_rows(values_only=True)]


def _extract_team_tracker_format(path: Path) -> list:
    """
    THIS IS THE ACTUAL 2b LIMITLESS LAYOUT (confirmed from their live tracker):
    one column per CliftonStrengths theme, grouped visually under domain banners
    (EXECUTING / INFLUENCING / RELATIONSHIP BUILDING / STRATEGIC THINKING), with
    the theme names themselves living a few rows down — the tracker has THREE
    stacked header rows (domain banner, domain description, theme name) before
    the data starts, and often a totals row at the bottom. This scans for
    whichever row actually contains the 34 theme names, wherever it is, and
    ignores rows that don't look like a real person (blank name, totals row).
    """
    grid = _load_raw_grid(path)

    header_row_idx, best_count = None, 0
    for i in range(min(10, len(grid))):
        count = sum(1 for cell in grid[i] if _normalise_theme(cell))
        if count > best_count:
            header_row_idx, best_count = i, count

    if header_row_idx is None or best_count < 25:
        return []  # doesn't look like this format — let the caller try something else

    header_row = grid[header_row_idx]
    theme_cols = [(idx, _normalise_theme(v)) for idx, v in enumerate(header_row) if _normalise_theme(v)]
    theme_col_idxs = {idx for idx, _ in theme_cols}

    name_idx = None
    for idx, v in enumerate(header_row):
        if isinstance(v, str) and re.search(r'name|team member|employee', v, re.IGNORECASE):
            name_idx = idx
            break
    if name_idx is None:
        left_of_themes = [idx for idx in range(len(header_row)) if idx not in theme_col_idxs]
        name_idx = left_of_themes[0] if left_of_themes else 0

    reports = []
    for row in grid[header_row_idx + 1:]:
        if name_idx >= len(row):
            continue
        name_val = row[name_idx]
        if not isinstance(name_val, str) or not name_val.strip():
            continue  # skips blank rows and the numeric totals row at the bottom

        ranks = {}
        for idx, theme in theme_cols:
            if idx >= len(row) or row[idx] is None:
                continue
            try:
                rank_num = int(float(row[idx]))
            except (ValueError, TypeError):
                continue
            if 1 <= rank_num <= 34:
                ranks[rank_num] = theme

        if len(ranks) < 34:
            missing = [n for n in range(1, 35) if n not in ranks]
            print(f"WARNING '{name_val.strip()}' in {path.name}: missing rank(s) {missing} — row skipped",
                  file=sys.stderr)
            continue

        reports.append({"name": name_val.strip(), "date": "", "ranks": ranks, "source_file": path.name})

    return reports


def extract_from_table(path: Path) -> list:
    """Returns a LIST of report dicts — a single spreadsheet export may contain many people."""
    team_tracker_reports = _extract_team_tracker_format(path)
    if team_tracker_reports:
        return team_tracker_reports

    if path.suffix.lower() == ".csv":
        df = pd.read_csv(path)
    else:
        df = pd.read_excel(path)

    cols_lower = {c: str(c).strip().lower() for c in df.columns}

    def find_col(*keywords):
        for orig, low in cols_lower.items():
            if any(k in low for k in keywords):
                return orig
        return None

    name_col = find_col("name")
    theme_col = find_col("theme", "talent", "strength")
    rank_col = find_col("rank", "sequence", "position", "order")

    reports = []

    if name_col and theme_col and rank_col:
        # "Long" format: one row per (person, theme) pair.
        for person, group in df.groupby(name_col):
            group_sorted = group.sort_values(rank_col)
            ranks = {}
            for i, (_, r) in enumerate(group_sorted.iterrows(), start=1):
                theme = _normalise_theme(r[theme_col])
                rank_val = int(r[rank_col]) if str(r[rank_col]).strip().isdigit() else i
                if theme:
                    ranks[rank_val] = theme
            _require_complete(ranks, f"{path.name} ({person})")
            reports.append({"name": str(person).strip(), "date": "", "ranks": ranks, "source_file": path.name})
        return reports

    # "Wide" format: one row per person, 34 theme columns (e.g. "Rank 1".."Rank 34",
    # "1".."34", or "Theme 1".."Theme 34") — the cell holds the THEME NAME for that rank.
    rank_like_cols = []
    for c in df.columns:
        m = re.search(r'(\d{1,2})\s*$', str(c))
        if m and 1 <= int(m.group(1)) <= 34:
            rank_like_cols.append((int(m.group(1)), c))

    if len(rank_like_cols) >= 30:  # tolerate a couple missing/misnamed columns
        rank_like_cols.sort()
        for _, row in df.iterrows():
            person = str(row[name_col]).strip() if name_col else f"Person row {_ + 2}"
            ranks = {}
            for n, col in rank_like_cols:
                theme = _normalise_theme(row[col])
                if theme:
                    ranks[n] = theme
            _require_complete(ranks, f"{path.name} ({person})")
            reports.append({"name": person, "date": "", "ranks": ranks, "source_file": path.name})
        return reports

    # Bare single column of 34 theme names, in rank order, for ONE person.
    first_col = df.columns[0]
    values = [v for v in df[first_col].tolist() if pd.notna(v)]
    ranks = {}
    for i, v in enumerate(values, start=1):
        theme = _normalise_theme(v)
        if theme:
            ranks[i] = theme
        if i > 34:
            break
    _require_complete(ranks, path.name)
    reports.append({"name": path.stem.replace("_", " ").replace("-", " "), "date": "",
                     "ranks": ranks, "source_file": path.name})
    return reports


def _require_complete(ranks: dict, source: str):
    missing = [n for n in range(1, 35) if n not in ranks]
    if missing:
        raise ValueError(
            f"{source}: could not identify rank(s) {missing} out of 34 — "
            "layout may differ from what this parser expects. Send a sample so the parser can be adjusted."
        )


def leading_domain(ranks: dict) -> str:
    """Matches the Gallup report's own 'You lead with X themes' — domain of the #1 theme."""
    return DOMAIN_MAP[ranks[1]]


def domain_counts_top10(ranks: dict) -> dict:
    counts = {d: 0 for d in DOMAIN_ORDER}
    for n in range(1, 11):
        counts[DOMAIN_MAP[ranks[n]]] += 1
    return counts


# ---------------------------------------------------------------------------
# Domain "wheel" chart — one quadrant per domain, spokes = that domain's
# themes, spoke length = strength (rank 1 reaches the rim, rank 34 sits at
# the centre). Matches the style of wheel the team already builds by hand.
# ---------------------------------------------------------------------------
DOMAIN_CHART_COLORS = {
    "Executing": "#8064A2",              # purple
    "Influencing": "#F79646",            # orange
    "Relationship Building": "#4F81BD",  # blue
    "Strategic Thinking": "#9BBB59",     # green
}

# Light tint of each domain color — used as the background for EVERY cell in that
# domain's columns; the deep DOMAIN_CHART_COLORS above is reserved for a person's
# top-5 ranked cells only, so the highlight still stands out against the tint.
DOMAIN_LIGHT_HEX = {
    "Executing": "E2D5F0",
    "Influencing": "FCE4D6",
    "Relationship Building": "DDEBF7",
    "Strategic Thinking": "E2EFDA",
}


def _wheel_layout() -> list:
    """34 (domain, theme, angle_in_radians) entries, each domain filling a fixed 90° quadrant,
    starting at 12 o'clock and going clockwise: Executing, Influencing, Relationship Building,
    Strategic Thinking."""
    layout = []
    for q, domain in enumerate(DOMAIN_ORDER):
        themes = DOMAIN_THEMES_ORDERED[domain]
        n = len(themes)
        for i, theme in enumerate(themes):
            angle_deg = q * 90 + (i + 0.5) / n * 90
            layout.append((domain, theme, np.radians(angle_deg)))
    return layout


_WHEEL_LAYOUT = _wheel_layout()


def generate_wheel_png(name: str, rank_of_theme: dict, out_path: Path):
    """rank_of_theme: {theme_name: rank_number (1-34)}. Saves a PNG to out_path."""
    fig = plt.figure(figsize=(3.6, 3.6), dpi=150)
    ax = fig.add_subplot(111, projection="polar")
    ax.set_theta_zero_location("N")
    ax.set_theta_direction(-1)
    ax.set_ylim(0, 34)
    ax.set_yticklabels([])
    ax.grid(color="#DDDDDD", linewidth=0.5)
    ax.spines["polar"].set_visible(False)

    by_domain = {}
    for domain, theme, angle in _WHEEL_LAYOUT:
        by_domain.setdefault(domain, []).append((theme, angle))

    for q, domain in enumerate(DOMAIN_ORDER):
        entries = by_domain[domain]
        quadrant_start = np.radians(q * 90)
        quadrant_end = np.radians((q + 1) * 90)
        # wedge: start at centre along the quadrant's lower boundary, out to each
        # spoke tip in turn, back to centre along the quadrant's upper boundary
        thetas = [quadrant_start] + [angle for _, angle in entries] + [quadrant_end]
        radii = [0.0] + [34 - rank_of_theme[theme] + 1 for theme, _ in entries] + [0.0]
        color = DOMAIN_CHART_COLORS[domain]
        ax.fill(thetas, radii, color=color, alpha=0.55, linewidth=0.8, edgecolor=color)

    tick_angles = [angle for _, _, angle in _WHEEL_LAYOUT]
    tick_labels = [theme for _, theme, _ in _WHEEL_LAYOUT]
    ax.set_xticks(tick_angles)
    ax.set_xticklabels(tick_labels, fontsize=4.3)
    ax.tick_params(axis="x", pad=2)

    fig.tight_layout(pad=0.3)
    fig.savefig(out_path, transparent=True)
    plt.close(fig)


def top_n_themes(rank_of_theme: dict, n: int = 5) -> list:
    """rank_of_theme: {theme_name: rank_number}. Returns the n theme names with the
    lowest rank number (1 = strongest), in rank order."""
    return [theme for theme, _rank in sorted(rank_of_theme.items(), key=lambda kv: kv[1])[:n]]


def _render_card_face(name: str, top5: list, out_path: Path, size_in=(8.5, 4.25), dpi=150):
    """Renders ONE card face (name + top 5 strengths) as a standalone PNG."""
    fig = plt.figure(figsize=size_in, dpi=dpi)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.axis("off")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)

    ax.text(0.5, 0.82, name.title(), fontsize=28, fontweight="bold",
            ha="center", va="center", family="sans-serif")
    ax.text(0.5, 0.68, "TOP 5 STRENGTHS", fontsize=11, color="#888888",
            ha="center", va="center", family="sans-serif")

    y = 0.52
    for theme in top5:
        domain = DOMAIN_MAP[theme]
        color = DOMAIN_CHART_COLORS[domain]
        ax.add_patch(plt.Rectangle((0.30, y - 0.028), 0.03, 0.056, color=color))
        ax.text(0.37, y, theme, fontsize=17, ha="left", va="center", family="sans-serif")
        y -= 0.115

    fig.savefig(out_path, dpi=dpi)
    plt.close(fig)


def generate_name_card_png(name: str, top5: list, out_path: Path):
    """Renders a full US-Letter page (portrait) with the SAME card content printed
    twice: right-side-up in the bottom half, upside-down in the top half — so when
    the page is folded in half at the middle, it reads correctly from both sides as
    a table tent. top5: list of theme names, strongest first."""
    from PIL import Image, ImageDraw

    dpi = 150
    page_w, page_h = int(8.5 * dpi), int(11 * dpi)
    half_w, half_h = page_w, page_h // 2

    with tempfile.TemporaryDirectory() as tmp:
        face_path = Path(tmp) / "face.png"
        _render_card_face(name, top5, face_path, size_in=(8.5, 4.25), dpi=dpi)
        face = Image.open(face_path).convert("RGBA")
        face = face.resize((half_w, half_h))
        face_rotated = face.rotate(180)

        page = Image.new("RGBA", (page_w, page_h), "white")
        page.paste(face_rotated, (0, 0), face_rotated)
        page.paste(face, (0, page_h - half_h), face)

        draw = ImageDraw.Draw(page)
        y_fold = page_h // 2
        dash_len, gap_len, x = 10, 8, int(0.05 * page_w)
        while x < page_w - int(0.05 * page_w):
            draw.line([(x, y_fold), (min(x + dash_len, page_w), y_fold)], fill="#BBBBBB", width=2)
            x += dash_len + gap_len

        page.convert("RGB").save(out_path)


def _set_cell_shading(cell, color_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _set_landscape(document):
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Inches(0.4)
    section.top_margin = section.bottom_margin = Inches(0.4)
    return section


def _set_tabloid_landscape(document):
    """Workshop grid has 34 narrow theme columns plus name/domain columns — that
    doesn't fit a Letter-size page at a readable font, so this document uses
    Tabloid (11x17in) landscape instead, which is what a print shop / office
    printer set to 'ledger' or 'tabloid' paper will produce."""
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(17)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(0.4)
    section.top_margin = section.bottom_margin = Inches(0.4)
    return section


def _set_col_widths(table, widths):
    """Force explicit column widths on every cell AND on the underlying <w:tblGrid>
    with a FIXED table layout. Setting cell.width alone is silently ignored by
    Word/LibreOffice's autofit, which redistributes columns roughly evenly and
    causes header text to wrap letter-by-letter in a narrow-name-name column —
    the failure mode from the very first version of this table."""
    table.autofit = False

    tbl = table._tbl
    tblPr = tbl.tblPr
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")

    tblGrid = tbl.find(qn("w:tblGrid"))
    for gridCol, width in zip(tblGrid.findall(qn("w:gridCol")), widths):
        gridCol.set(qn("w:w"), str(width.twips))

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width


def _add_full_grid_table(document, people_rows: list):
    """ONE continuous table (Name + Leading Domain + all 34 theme columns), not
    split by domain — matches the team's own tracker, which is one wide sheet.
    Report Date / Source File are dropped from this printed view (they're
    bookkeeping columns, not something a workshop room needs to see); they're
    still in the data spreadsheet. Column widths are computed to exactly fill a
    Tabloid-landscape page and enforced via a fixed table layout so headers wrap
    normally (word-by-word) instead of collapsing to letter-by-letter."""
    info_cols = ["Name", "Leading Domain"]
    usable_width = Inches(17) - 2 * Inches(0.4)
    name_width = Inches(1.3)
    domain_width = Inches(1.0)
    theme_col_width = Emu(int((usable_width - name_width - domain_width) / len(ORDERED_THEMES)))
    info_widths = [name_width, domain_width]
    col_widths = info_widths + [theme_col_width] * len(ORDERED_THEMES)

    table = document.add_table(rows=2 + len(people_rows), cols=len(info_cols) + len(ORDERED_THEMES))

    # Row 1: domain banner, merged across each domain's theme columns
    banner = table.rows[0].cells
    for c in banner[:len(info_cols)]:
        c.text = ""
    col = len(info_cols)
    for domain in DOMAIN_ORDER:
        span = len(DOMAIN_THEMES_ORDERED[domain])
        merged_cell = banner[col]
        for extra in range(1, span):
            merged_cell = merged_cell.merge(banner[col + extra])
        merged_cell.text = domain.upper()
        _set_cell_shading(merged_cell, DOMAIN_CHART_COLORS[domain].lstrip("#"))
        merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in merged_cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        col += span

    # Row 2: column headers, all horizontal — theme names word-wrap onto 2 lines
    # in their narrow columns since the widths above are now actually enforced.
    header = table.rows[1].cells
    for i, h in enumerate(info_cols):
        header[i].text = h
        _set_cell_shading(header[i], "305496")
        header[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, theme in enumerate(ORDERED_THEMES, start=len(info_cols)):
        cell = header[i]
        cell.text = theme
        domain = DOMAIN_MAP[theme]
        _set_cell_shading(cell, DOMAIN_CHART_COLORS[domain].lstrip("#"))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(6.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    table.rows[1].height = Inches(0.55)

    # Data rows
    for r, (name, row) in enumerate(people_rows, start=2):
        rank_of_theme = {t: row[t] for t in ORDERED_THEMES}
        top5 = set(top_n_themes(rank_of_theme, 5))

        cells = table.rows[r].cells
        cells[0].text = name
        cells[1].text = str(row.get("Leading Domain", ""))
        for c in cells[:len(info_cols)]:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(7)

        for i, theme in enumerate(ORDERED_THEMES, start=len(info_cols)):
            cell = cells[i]
            cell.text = str(row[theme])
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            domain = DOMAIN_MAP[theme]
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(7)
            if theme in top5:
                _set_cell_shading(cell, DOMAIN_CHART_COLORS[domain].lstrip("#"))
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                # Neutral alternating column banding (matches the reference tracker,
                # which shades whole columns gray/white for readability — domain
                # color is reserved for the header banner and each person's top 5).
                col_index = i - len(info_cols)
                _set_cell_shading(cell, "F2F2F2" if col_index % 2 == 0 else "FFFFFF")

    _set_col_widths(table, col_widths)
    return table


def build_workshop_docx(merged: dict, output_path: Path):
    """Builds the Word 'workshop workbook' — one continuous team rank table (all
    34 themes, no domain-split tables) on a Tabloid-landscape page, followed by
    every person's wheel chart, 2 per row, with a page break after each row so a
    chart never gets sliced across a page. Pure python-docx — no external runtime
    needed, so this can run inside the deployed Streamlit app as well as the CLI."""
    document = Document()
    _set_tabloid_landscape(document)

    document.add_heading("CliftonStrengths Workshop Workbook", level=1)
    document.add_paragraph("Full team rank table — all 34 themes").italic = True

    people_rows = sorted(merged.items())
    _add_full_grid_table(document, people_rows)
    document.add_paragraph("")

    document.add_page_break()
    document.add_heading("Strengths Wheels", level=1)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        for i in range(0, len(people_rows), 2):
            pair = people_rows[i:i + 2]
            table = document.add_table(rows=2, cols=2)
            table.autofit = True
            for col, (name, row) in enumerate(pair):
                rank_of_theme = {theme: row[theme] for theme in ORDERED_THEMES}
                png_path = tmp_dir / f"wheel_{i}_{col}.png"
                generate_wheel_png(name, rank_of_theme, png_path)

                name_cell = table.rows[0].cells[col]
                name_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = name_cell.paragraphs[0].add_run(name)
                run.bold = True

                img_cell = table.rows[1].cells[col]
                img_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_cell.paragraphs[0].add_run().add_picture(str(png_path), width=Inches(3.3))

            if i + 2 < len(people_rows):
                document.add_page_break()

        document.save(output_path)


def build_name_cards_docx(merged: dict, output_path: Path):
    """Builds the table name cards — one tent-fold card per person (name + top 5
    strengths, printed twice so folding the page in half reads correctly from both
    sides), each on its own page."""
    document = Document()
    section = document.sections[0]
    MARGIN = Inches(0.1)
    section.left_margin = section.right_margin = MARGIN
    section.top_margin = section.bottom_margin = MARGIN
    # shrink slightly below the true printable area — some renderers add a hair of
    # line-height around the image's paragraph, which otherwise spills a blank page
    SAFETY = 0.96
    printable_w = int((section.page_width - 2 * MARGIN) * SAFETY)
    printable_h = int((section.page_height - 2 * MARGIN) * SAFETY)

    people_rows = sorted(merged.items())
    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        for i, (name, row) in enumerate(people_rows):
            rank_of_theme = {theme: row[theme] for theme in ORDERED_THEMES}
            top5 = top_n_themes(rank_of_theme, 5)
            png_path = tmp_dir / f"card_{i}.png"
            generate_name_card_png(name, top5, png_path)

            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            # explicit width AND height, sized to the printable area exactly — the
            # source PNG is already an 8.5x11 page, so this just fits it to the
            # section's printable box regardless of any DPI metadata on the file.
            paragraph.add_run().add_picture(str(png_path), width=printable_w, height=printable_h)

            if i < len(people_rows) - 1:
                document.add_page_break()

        document.save(output_path)


def build_wheels_workbook(merged: dict, output_path: Path):
    """Builds a SEPARATE workbook containing just the 'Strengths Wheels' chart grid —
    2 people per row, each with a name banner above their wheel. A manual page break
    is inserted after every row of 2 charts so a wheel never gets sliced across a
    printed/exported page."""
    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        wb = Workbook()
        ws = wb.active
        ws.title = "Strengths Wheels"
        ws.page_setup.orientation = "landscape"
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 0
        ws.sheet_properties.pageSetUpPr.fitToPage = True

        banner_fill = PatternFill(start_color="BDD7EE", end_color="BDD7EE", fill_type="solid")
        banner_font = Font(name="Arial", bold=True, size=12)

        CHART_COLS = 7       # columns each chart+banner spans
        CHART_ROW_SPAN = 20  # rows each chart occupies below its banner
        GAP_ROWS = 2         # blank rows between one chart-row and the next
        GAP_COLS = 1
        BLOCK_ROWS = 1 + CHART_ROW_SPAN + GAP_ROWS  # banner + chart + gap

        for col in range(1, 2 * CHART_COLS + GAP_COLS + 2):
            ws.column_dimensions[get_column_letter(col)].width = 9

        people = sorted(merged.items())
        last_grid_row = -1
        for i, (name, row) in enumerate(people):
            rank_of_theme = {theme: row[theme] for theme in ORDERED_THEMES}

            png_path = tmp_dir / f"wheel_{i}.png"
            generate_wheel_png(name, rank_of_theme, png_path)

            grid_row, grid_col = divmod(i, 2)
            top_row = grid_row * BLOCK_ROWS + 1
            left_col = grid_col * (CHART_COLS + GAP_COLS) + 1

            # start a fresh printed page for every new row of charts
            if grid_row != last_grid_row and grid_row > 0:
                ws.row_breaks.append(Break(id=top_row - 1))
            last_grid_row = grid_row

            ws.merge_cells(start_row=top_row, start_column=left_col,
                            end_row=top_row, end_column=left_col + CHART_COLS - 1)
            banner_cell = ws.cell(row=top_row, column=left_col, value=name)
            banner_cell.fill = banner_fill
            banner_cell.font = banner_font
            banner_cell.alignment = Alignment(horizontal="center", vertical="center")
            ws.row_dimensions[top_row].height = 22

            img = XLImage(str(png_path))
            img.width = 330
            img.height = 330
            ws.add_image(img, f"{get_column_letter(left_col)}{top_row + 1}")

        wb.save(output_path)
    return len(people)


def collect_reports(input_path: Path) -> list:
    if input_path.is_dir():
        files = sorted(
            [p for p in input_path.iterdir()
             if p.suffix.lower() in (".pdf", ".txt", ".csv", ".xlsx", ".xls")]
        )
    else:
        files = [input_path]

    if not files:
        print(f"No supported files (.pdf/.txt/.csv/.xlsx) found at {input_path}", file=sys.stderr)
        sys.exit(1)

    reports = []
    for path in files:
        try:
            suffix = path.suffix.lower()
            if suffix == ".pdf":
                reports.append(extract_from_pdf(path))
            elif suffix == ".txt":
                reports.append(extract_from_text(path))
            elif suffix in (".csv", ".xlsx", ".xls"):
                reports.extend(extract_from_table(path))
            print(f"OK: {path.name}")
        except Exception as exc:  # noqa: BLE001 - surface per-file errors, keep going
            print(f"SKIPPED {path.name}: {exc}", file=sys.stderr)
    return reports


# ---------------------------------------------------------------------------
# Spreadsheet output
# ---------------------------------------------------------------------------
HEADER_FONT = Font(name="Arial", bold=True, color="FFFFFF")
HEADER_FILL = PatternFill(start_color="305496", end_color="305496", fill_type="solid")
BODY_FONT = Font(name="Arial")
DOMAIN_FILL = {d: PatternFill(start_color=hex_, end_color=hex_, fill_type="solid")
               for d, hex_ in DOMAIN_LIGHT_HEX.items()}


DOMAIN_BANNER_FILL = {d: PatternFill(start_color=DOMAIN_CHART_COLORS[d].lstrip("#"),
                                      end_color=DOMAIN_CHART_COLORS[d].lstrip("#"),
                                      fill_type="solid")
                       for d in DOMAIN_ORDER}

# Same deep domain colors as the banners — used to highlight only a person's TOP 5
# ranked themes (rank 1-5), matching the team's own tracker: everything else stays
# plain white, so the five deep-colored cells are what draws the eye.
DOMAIN_TOP5_FILL = DOMAIN_BANNER_FILL
TOP5_FONT = Font(name="Arial", bold=True, color="FFFFFF")
INFO_COLUMNS = ["Name", "Report Date", "Source File", "Leading Domain"]

# Neutral alternating column banding for non-top5 cells, matching the reference
# tracker (whole columns shaded gray/white for readability; domain color is
# reserved for the banner row and each person's top-5 highlight).
BAND_FILL_GRAY = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
BAND_FILL_WHITE = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")


def build_workbook(reports: list, output_path: Path):
    """Builds the data workbook in the SAME layout as the team's own tracker:
    a domain-banner row, then a row of the 34 theme names as column headers,
    then one row per person with their RANK NUMBER (not the theme name) under
    each theme column."""
    existing_rows = {}
    if output_path.exists():
        wb_existing = load_workbook(output_path)
        if "CliftonStrengths Data" in wb_existing.sheetnames:
            ws_existing = wb_existing["CliftonStrengths Data"]
            headers = [c.value for c in ws_existing[2]]  # theme-name row is row 2
            for row in ws_existing.iter_rows(min_row=3, values_only=True):
                row_dict = dict(zip(headers, row))
                if row_dict.get("Name"):
                    existing_rows[row_dict["Name"]] = row_dict

    wb = Workbook()
    ws = wb.active
    ws.title = "CliftonStrengths Data"

    all_headers = INFO_COLUMNS + ORDERED_THEMES

    # Row 1: domain banner, merged across each domain's theme columns
    col = len(INFO_COLUMNS) + 1
    for domain in DOMAIN_ORDER:
        span = len(DOMAIN_THEMES_ORDERED[domain])
        ws.merge_cells(start_row=1, start_column=col, end_row=1, end_column=col + span - 1)
        cell = ws.cell(row=1, column=col, value=domain.upper())
        cell.font = Font(name="Arial", bold=True, color="FFFFFF")
        cell.fill = DOMAIN_BANNER_FILL[domain]
        cell.alignment = Alignment(horizontal="center")
        col += span
    ws.row_dimensions[1].height = 20

    # Row 2: actual column headers (theme names)
    for col_idx, h in enumerate(all_headers, start=1):
        cell = ws.cell(row=2, column=col_idx, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
    ws.freeze_panes = "E3"
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    merged = dict(existing_rows)
    for r in reports:
        row = {"Name": r["name"], "Report Date": r["date"], "Source File": r["source_file"],
               "Leading Domain": leading_domain(r["ranks"])}
        for rank_num, theme in r["ranks"].items():
            row[theme] = rank_num
        merged[r["name"]] = row

    for row_idx, (name, row) in enumerate(sorted(merged.items()), start=3):
        rank_of_theme = {theme: row[theme] for theme in ORDERED_THEMES}
        top5 = set(top_n_themes(rank_of_theme, 5))
        for col_idx, h in enumerate(all_headers, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=row.get(h, ""))
            cell.alignment = Alignment(horizontal="center") if h in DOMAIN_MAP else Alignment()
            domain = DOMAIN_MAP.get(h)
            if domain and h in top5:
                cell.fill = DOMAIN_TOP5_FILL[domain]
                cell.font = TOP5_FONT
            elif domain:
                theme_col_index = ORDERED_THEMES.index(h)
                cell.fill = BAND_FILL_GRAY if theme_col_index % 2 == 0 else BAND_FILL_WHITE
                cell.font = BODY_FONT
            else:
                cell.font = BODY_FONT

    for col_idx, h in enumerate(all_headers, start=1):
        width = 9 if h in DOMAIN_MAP else max(16, len(h) + 2)
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    wb.save(output_path)
    return merged


def main():
    parser = argparse.ArgumentParser(description="Extract CliftonStrengths 34 results (PDF/Excel/CSV/pasted text) into a spreadsheet, a wheel-chart workbook, a workshop Word workbook, and printable name cards.")
    parser.add_argument("--input", required=True, help="A single file or a folder mixing .pdf/.txt/.csv/.xlsx files.")
    parser.add_argument("--output", required=True, help="Path to the output data .xlsx file.")
    parser.add_argument("--charts-output", default=None,
                         help="Path to the SEPARATE wheel-chart .xlsx file. Defaults to '<output>_Wheels.xlsx'.")
    parser.add_argument("--no-charts", action="store_true",
                         help="Skip generating the wheel-chart .xlsx file.")
    parser.add_argument("--workbook-output", default=None,
                         help="Path to the workshop Word workbook (.docx). Defaults to '<output>_Workshop_Workbook.docx'.")
    parser.add_argument("--no-workbook", action="store_true",
                         help="Skip generating the workshop Word workbook.")
    parser.add_argument("--cards-output", default=None,
                         help="Path to the table name cards (.docx). Defaults to '<output>_Name_Cards.docx'.")
    parser.add_argument("--no-cards", action="store_true",
                         help="Skip generating the table name cards.")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    charts_output = Path(args.charts_output) if args.charts_output else \
        output_path.with_name(output_path.stem + "_Wheels" + output_path.suffix)
    workbook_output = Path(args.workbook_output) if args.workbook_output else \
        output_path.with_name(output_path.stem + "_Workshop_Workbook.docx")
    cards_output = Path(args.cards_output) if args.cards_output else \
        output_path.with_name(output_path.stem + "_Name_Cards.docx")

    reports = collect_reports(input_path)
    print(f"Parsed {len(reports)} report(s) successfully.")

    if not reports and not output_path.exists():
        print("Nothing parsed — no output file exists to merge into either.", file=sys.stderr)
        return

    merged = build_workbook(reports, output_path)
    print(f"Wrote {len(merged)} total row(s) to {output_path}")

    if not args.no_charts:
        chart_count = build_wheels_workbook(merged, charts_output)
        print(f"Wrote {chart_count} wheel chart(s) to {charts_output}")

    if not args.no_workbook:
        build_workshop_docx(merged, workbook_output)
        print(f"Wrote workshop workbook to {workbook_output}")

    if not args.no_cards:
        build_name_cards_docx(merged, cards_output)
        print(f"Wrote {len(merged)} name card(s) to {cards_output}")


if __name__ == "__main__":
    main()
